#!/usr/bin/env python3
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
from lxml import html, etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import sys
import tempfile


if len(sys.argv) < 3:
    print("Usage: export-paper-docx-fallback.py report.clean.html report.docx", file=sys.stderr)
    sys.exit(2)

clean_html = Path(sys.argv[1])
output_docx = Path(sys.argv[2])

ACCENT = "0072B2"
ACCENT2 = "009E73"
RULE = "B7D7D0"
SOFT = "F3F8F6"
CALLOUT = "F1F8F5"
GRID = "DCE6E3"
CHART_COLORS = ["0072B2", "009E73", "E69F00", "CC79A7", "56B4E9", "D55E00"]
CHART_MARKER = "[[NATIVE_CHART_GBIF_RECORDS]]"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
C_NS = "http://schemas.openxmlformats.org/drawingml/2006/chart"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

NS = {
    "w": W_NS,
    "r": R_NS,
    "wp": WP_NS,
    "a": A_NS,
    "c": C_NS,
    "rel": REL_NS,
    "ct": CT_NS,
}


def W(tag):
    return qn(f"w:{tag}")


def Q(prefix, tag):
    return f"{{{NS[prefix]}}}{tag}"


def text_content(el):
    return " ".join(el.text_content().split())


def parse_number(value):
    cleaned = re.sub(r"[^0-9.\-]", "", value or "")
    if not cleaned:
        return 0
    return float(cleaned) if "." in cleaned else int(cleaned)


def set_run_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(W("eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text_runs(paragraph, element, size=10.5, bold=False, italic=False):
    tag = element.tag.lower() if isinstance(element.tag, str) else ""
    local_bold = bold or tag in {"strong", "b"}
    local_italic = italic or tag in {"em", "i"}
    if element.text:
        run = paragraph.add_run(element.text)
        set_run_font(run, size, local_bold, local_italic)
    for child in element:
        add_text_runs(paragraph, child, size, local_bold, local_italic)
        if child.tail:
            run = paragraph.add_run(child.tail)
            set_run_font(run, size, bold, italic)


def paragraph_border(paragraph, **borders):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(W("pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side, opts in borders.items():
        old = p_bdr.find(W(side))
        if old is not None:
            p_bdr.remove(old)
        el = OxmlElement(f"w:{side}")
        for key, value in opts.items():
            el.set(W(key), str(value))
        p_bdr.append(el)


def add_plain(document, text, style=None, size=10.5, bold=False, color=None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.45
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, color=color)
    return paragraph


def add_para_from_el(document, element, style=None, size=10.5, color=None, bold=False):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.line_spacing = 1.45
    paragraph.paragraph_format.space_after = Pt(6)
    add_text_runs(paragraph, element, size=size, bold=bold)
    if color:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def cell_shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(W("shd"))
    if old is not None:
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(W("fill"), fill)
    tc_pr.append(shd)


def cell_borders(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(W("tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side, opts in kwargs.items():
        old = borders.find(W(side))
        if old is not None:
            borders.remove(old)
        el = OxmlElement(f"w:{side}")
        for key, value in opts.items():
            el.set(W(key), str(value))
        borders.append(el)


def cell_margins(cell, top=95, start=115, bottom=95, end=115):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(W("tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        old = mar.find(W(side))
        if old is not None:
            mar.remove(old)
        el = OxmlElement(f"w:{side}")
        el.set(W("w"), str(val))
        el.set(W("type"), "dxa")
        mar.append(el)


def set_cell_text(cell, text, size=9.2, bold=False, italic=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.22
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold, italic, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_margins(cell)


def fixed_table_grid(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell.width = Inches(widths[index])


def add_docx_table(document, table_el):
    rows = table_el.xpath("./thead/tr | ./tbody/tr | ./tr")
    if not rows:
        return
    col_count = max(len(row.xpath("./th|./td")) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_text = " | ".join(text_content(cell) for cell in rows[0].xpath("./th|./td"))
    if header_text.startswith("常用英文名"):
        widths = [1.25, 1.25, 2.55, 2.35]
    elif header_text.startswith("学名") and "GBIF" in header_text:
        widths = [1.1, 0.9, 2.65, 2.75]
    else:
        widths = [6.8 / col_count] * col_count
    fixed_table_grid(table, widths)

    for i, row_el in enumerate(rows):
        row = table.rows[i]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(28 if i else 31)
        cells = row_el.xpath("./th|./td")
        is_header = bool(row_el.xpath("./th")) or (
            i == 0 and row_el.getparent() is not None and row_el.getparent().tag.lower() == "thead"
        )
        for j in range(col_count):
            cell = table.cell(i, j)
            text = text_content(cells[j]) if j < len(cells) else ""
            border = {"val": "single", "sz": "4", "space": "0", "color": GRID}
            cell_borders(cell, top=border, bottom=border, left=border, right=border)
            if is_header:
                cell_shade(cell, ACCENT)
                set_cell_text(cell, text, size=9.4, bold=True, color="FFFFFF")
            else:
                if i % 2 == 0:
                    cell_shade(cell, SOFT)
                set_cell_text(cell, text, size=9.15, italic=(j == 0 and text.startswith("Lepus")))
    document.add_paragraph()


def add_callout(document, element):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.85)
    cell_shade(cell, CALLOUT)
    cell_margins(cell, top=180, start=220, bottom=180, end=180)
    cell_borders(
        cell,
        left={"val": "single", "sz": "24", "space": "0", "color": ACCENT2},
        top={"val": "nil"},
        bottom={"val": "nil"},
        right={"val": "nil"},
    )
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.45
    paragraph.paragraph_format.space_after = Pt(0)
    add_text_runs(paragraph, element, size=10.5)
    document.add_paragraph()


def collect_chart_data(root):
    rows = root.xpath('//*[contains(concat(" ", normalize-space(@class), " "), " bar-row ")]')
    labels = []
    values = []
    for row in rows:
        spans = row.xpath(".//span")
        if len(spans) >= 2:
            labels.append(text_content(spans[0]))
            values.append(parse_number(text_content(spans[-1])))
    if labels and values:
        return {"title": "主要物种的公开记录规模", "labels": labels, "values": values}
    return None


def build_document(clean_html_path, output_path):
    root = html.fromstring(Path(clean_html_path).read_text(encoding="utf-8"))
    body = root.find("body")
    chart_data = collect_chart_data(root)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.76)
    section.right_margin = Inches(0.76)

    styles = document.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(W("eastAsia"), "宋体")
        style.font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.45
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 12)]:
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor.from_string(ACCENT)

    source_counter = {"value": 0}

    def handle(element):
        tag = element.tag.lower() if isinstance(element.tag, str) else ""
        classes = element.get("class", "")
        if tag in {"section", "main", "div"}:
            if "callout" in classes:
                add_callout(document, element)
                return
            if "subtitle" in classes:
                paragraph = add_para_from_el(document, element, size=14, bold=True, color=ACCENT)
                paragraph_border(paragraph, bottom={"val": "single", "sz": "6", "space": "6", "color": RULE})
                return
            if "bar-row" in classes:
                return
            for child in element:
                handle(child)
            return
        if tag == "h1":
            paragraph = add_para_from_el(document, element, style="Title", size=20, bold=True, color=ACCENT)
            paragraph_border(paragraph, left={"val": "single", "sz": "18", "space": "6", "color": ACCENT})
            return
        if tag == "h2":
            paragraph = add_para_from_el(document, element, style="Heading 1", size=16, bold=True, color=ACCENT)
            paragraph_border(paragraph, bottom={"val": "single", "sz": "6", "space": "4", "color": RULE})
            return
        if tag == "h3":
            title = text_content(element)
            add_para_from_el(document, element, style="Heading 2", size=12.5, bold=True, color=ACCENT)
            if chart_data and title.startswith("可视化"):
                add_plain(document, CHART_MARKER)
            return
        if tag == "p":
            text = text_content(element)
            if not text:
                return
            if "source-item" in classes:
                source_counter["value"] += 1
                add_plain(document, f'{source_counter["value"]}. {text}', size=8.8, color="555555")
                return
            if "caption" in classes:
                add_plain(document, text, size=8.8, color="666666")
                return
            add_para_from_el(document, element)
            return
        if tag in {"ol", "ul"}:
            style = "List Number" if tag == "ol" else "List Bullet"
            for li in element.xpath("./li"):
                add_plain(document, text_content(li), style=style)
            return
        if tag == "table":
            add_docx_table(document, element)
            return
        if tag == "hr":
            return
        text = text_content(element)
        if text:
            add_plain(document, text)

    if body is not None:
        for child in body:
            handle(child)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(root.xpath("string(//title)") or Path(output_path).stem)
    set_run_font(run, size=9, color="777777")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    postprocess_docx(Path(output_path), chart_data)


def text_of(paragraph):
    return "".join(paragraph.xpath(".//w:t/text()", namespaces=NS))


def next_rid(rels_root):
    max_id = 0
    for rel in rels_root.findall(Q("rel", "Relationship")):
        rid = rel.get("Id", "")
        if rid.startswith("rId") and rid[3:].isdigit():
            max_id = max(max_id, int(rid[3:]))
    return f"rId{max_id + 1}"


def make_chart_paragraph(rid):
    xml = f"""
    <w:p xmlns:w="{W_NS}" xmlns:r="{R_NS}" xmlns:wp="{WP_NS}" xmlns:a="{A_NS}" xmlns:c="{C_NS}">
      <w:pPr>
        <w:spacing w:before="120" w:after="120"/>
        <w:jc w:val="center"/>
      </w:pPr>
      <w:r>
        <w:drawing>
          <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="5486400" cy="2743200"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:docPr id="42" name="Native bar chart"/>
            <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
            <a:graphic>
              <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">
                <c:chart r:id="{rid}"/>
              </a:graphicData>
            </a:graphic>
          </wp:inline>
        </w:drawing>
      </w:r>
    </w:p>
    """
    return etree.fromstring(xml.encode("utf-8"))


def make_chart_xml(chart_data):
    labels = chart_data["labels"]
    values = chart_data["values"]
    cats = "\n".join(f'<c:pt idx="{i}"><c:v>{label}</c:v></c:pt>' for i, label in enumerate(labels))
    vals = "\n".join(f'<c:pt idx="{i}"><c:v>{value}</c:v></c:pt>' for i, value in enumerate(values))
    dpts = "\n".join(
        f"""
        <c:dPt>
          <c:idx val="{i}"/>
          <c:spPr>
            <a:solidFill><a:srgbClr val="{CHART_COLORS[i % len(CHART_COLORS)]}"/></a:solidFill>
            <a:ln><a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill></a:ln>
          </c:spPr>
        </c:dPt>
        """
        for i in range(len(labels))
    )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<c:chartSpace xmlns:c="{C_NS}" xmlns:a="{A_NS}" xmlns:r="{R_NS}">
  <c:date1904 val="0"/>
  <c:lang val="zh-CN"/>
  <c:roundedCorners val="0"/>
  <c:chart>
    <c:title>
      <c:tx><c:rich><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="zh-CN" sz="1200" b="1"/><a:t>{chart_data["title"]}</a:t></a:r></a:p></c:rich></c:tx>
      <c:overlay val="0"/>
    </c:title>
    <c:plotArea>
      <c:layout/>
      <c:barChart>
        <c:barDir val="bar"/>
        <c:grouping val="clustered"/>
        <c:varyColors val="1"/>
        <c:ser>
          <c:idx val="0"/>
          <c:order val="0"/>
          <c:tx><c:v>Records</c:v></c:tx>
          {dpts}
          <c:cat><c:strLit><c:ptCount val="{len(labels)}"/>{cats}</c:strLit></c:cat>
          <c:val><c:numLit><c:formatCode>#,##0</c:formatCode><c:ptCount val="{len(values)}"/>{vals}</c:numLit></c:val>
        </c:ser>
        <c:dLbls><c:showVal val="1"/><c:showLegendKey val="0"/><c:showCatName val="0"/><c:showSerName val="0"/><c:showPercent val="0"/><c:showBubbleSize val="0"/></c:dLbls>
        <c:axId val="12345678"/><c:axId val="87654321"/>
      </c:barChart>
      <c:catAx><c:axId val="12345678"/><c:scaling><c:orientation val="minMax"/></c:scaling><c:delete val="0"/><c:axPos val="l"/><c:majorTickMark val="none"/><c:minorTickMark val="none"/><c:tickLblPos val="nextTo"/><c:crossAx val="87654321"/><c:crosses val="autoZero"/><c:auto val="1"/><c:lblAlgn val="ctr"/><c:lblOffset val="100"/></c:catAx>
      <c:valAx><c:axId val="87654321"/><c:scaling><c:orientation val="minMax"/></c:scaling><c:delete val="0"/><c:axPos val="b"/><c:numFmt formatCode="#,##0" sourceLinked="0"/><c:majorGridlines/><c:majorTickMark val="out"/><c:minorTickMark val="none"/><c:tickLblPos val="nextTo"/><c:crossAx val="12345678"/><c:crosses val="autoZero"/><c:crossBetween val="between"/></c:valAx>
    </c:plotArea>
    <c:legend><c:legendPos val="b"/><c:overlay val="0"/></c:legend>
    <c:plotVisOnly val="1"/>
    <c:dispBlanksAs val="gap"/>
    <c:showDLblsOverMax val="0"/>
  </c:chart>
</c:chartSpace>
"""


def postprocess_docx(docx_path, chart_data):
    for prefix, uri in NS.items():
        if prefix not in {"rel", "ct"}:
            etree.register_namespace(prefix, uri)
    with tempfile.TemporaryDirectory() as tmp:
        work = Path(tmp)
        with ZipFile(docx_path) as zin:
            zin.extractall(work)
        document_path = work / "word/document.xml"
        rels_path = work / "word/_rels/document.xml.rels"
        types_path = work / "[Content_Types].xml"
        charts_dir = work / "word/charts"
        charts_dir.mkdir(exist_ok=True)

        document = etree.parse(str(document_path))
        body = document.getroot().find(Q("w", "body"))
        apply_fixed_table_widths(body)
        marker_para = None
        for paragraph in body.xpath(".//w:p", namespaces=NS):
            if text_of(paragraph) == CHART_MARKER:
                marker_para = paragraph
                break

        rels = etree.parse(str(rels_path))
        types = etree.parse(str(types_path))
        if marker_para is not None and chart_data:
            rid = next_rid(rels.getroot())
            rel = etree.SubElement(rels.getroot(), Q("rel", "Relationship"))
            rel.set("Id", rid)
            rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart")
            rel.set("Target", "charts/chart1.xml")

            parent = marker_para.getparent()
            index = parent.index(marker_para)
            parent.remove(marker_para)
            parent.insert(index, make_chart_paragraph(rid))

            if not any(el.get("PartName") == "/word/charts/chart1.xml" for el in types.getroot().findall(Q("ct", "Override"))):
                override = etree.SubElement(types.getroot(), Q("ct", "Override"))
                override.set("PartName", "/word/charts/chart1.xml")
                override.set("ContentType", "application/vnd.openxmlformats-officedocument.drawingml.chart+xml")
            (charts_dir / "chart1.xml").write_text(make_chart_xml(chart_data), encoding="utf-8")

        document.write(str(document_path), xml_declaration=True, encoding="UTF-8", standalone=True)
        rels.write(str(rels_path), xml_declaration=True, encoding="UTF-8", standalone=True)
        types.write(str(types_path), xml_declaration=True, encoding="UTF-8", standalone=True)

        rebuilt = work / "rebuilt.docx"
        with ZipFile(rebuilt, "w", ZIP_DEFLATED) as zout:
            for file in work.rglob("*"):
                if file.is_file() and file.name != "rebuilt.docx":
                    zout.write(file, file.relative_to(work).as_posix())
        rebuilt.replace(docx_path)


def first_row_text(table):
    return " | ".join(
        "".join(cell.xpath(".//w:t/text()", namespaces=NS))
        for cell in table.xpath("./w:tr[1]/w:tc", namespaces=NS)
    )


def apply_fixed_table_widths(body):
    for table in body.xpath(".//w:tbl", namespaces=NS):
        header = first_row_text(table)
        if header.startswith("一句话判断"):
            set_ooxml_table_width(table, [10052])
        elif header.startswith("常用英文名"):
            set_ooxml_table_width(table, [1650, 1650, 3350, 3070])
        elif header.startswith("学名 | GBIF"):
            set_ooxml_table_width(table, [1450, 1150, 3500, 3620])


def set_ooxml_table_width(table, widths):
    total = sum(widths)
    tbl_pr = table.find(Q("w", "tblPr"))
    if tbl_pr is None:
        tbl_pr = etree.Element(Q("w", "tblPr"))
        table.insert(0, tbl_pr)
    for tag in ["tblW", "tblLayout"]:
        old = tbl_pr.find(Q("w", tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = etree.Element(Q("w", "tblW"))
    tbl_w.set(Q("w", "w"), str(total))
    tbl_w.set(Q("w", "type"), "dxa")
    tbl_pr.insert(0, tbl_w)
    layout = etree.Element(Q("w", "tblLayout"))
    layout.set(Q("w", "type"), "fixed")
    tbl_pr.append(layout)

    old_grid = table.find(Q("w", "tblGrid"))
    if old_grid is not None:
        table.remove(old_grid)
    grid = etree.Element(Q("w", "tblGrid"))
    for width in widths:
        col = etree.Element(Q("w", "gridCol"))
        col.set(Q("w", "w"), str(width))
        grid.append(col)
    table.insert(list(table).index(tbl_pr) + 1, grid)

    for row in table.findall(Q("w", "tr")):
        cells = row.findall(Q("w", "tc"))
        for index, cell in enumerate(cells):
            width = widths[index] if index < len(widths) else widths[-1]
            tc_pr = cell.find(Q("w", "tcPr"))
            if tc_pr is None:
                tc_pr = etree.Element(Q("w", "tcPr"))
                cell.insert(0, tc_pr)
            old = tc_pr.find(Q("w", "tcW"))
            if old is not None:
                tc_pr.remove(old)
            tc_w = etree.Element(Q("w", "tcW"))
            tc_w.set(Q("w", "w"), str(width))
            tc_w.set(Q("w", "type"), "dxa")
            tc_pr.insert(0, tc_w)


build_document(clean_html, output_docx)
print(f"Wrote DOCX with python-docx fallback: {output_docx}")
